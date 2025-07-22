# resume_customizer_app/app.py
import streamlit as st
import docx
import io
import re
import openai
from docx.shared import Pt

# --- CONFIG ---
st.set_page_config(page_title="AI Resume Customizer", layout="wide")
st.title("🤖 AI Resume Tailor - ATS + Human Friendly")

# --- SIDEBAR INPUTS ---
openai_api_key = st.sidebar.text_input("🔑 Enter your OpenAI API Key", type="password")

with st.sidebar.expander("📝 Instructions"):
    st.markdown("""
    1. Upload your `.docx` resume  
    2. Paste any job description  
    3. Click 'Generate Custom Resume'  
    4. Download your new `.docx` resume (ATS-optimized + humanized)
    """)

# --- FILE UPLOAD ---
resume_file = st.file_uploader("📄 Upload your current resume (.docx)", type=["docx"])
jd_text = st.text_area("💼 Paste Job Description (JD)", height=300)

# --- ACTION BUTTON ---
if st.button("🚀 Generate Custom Resume"):
    if not resume_file or not jd_text or not openai_api_key:
        st.error("Please upload resume, paste JD, and enter OpenAI API key.")
    else:
        def extract_docx_text(file):
            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text), doc

        resume_text, doc = extract_docx_text(resume_file)

        prompt = f"""
        You are an expert resume writer. Take the following resume and job description, then:
        - Rewrite the summary using active, human, natural tone (no AI feel)
        - Inject all relevant JD keywords (skills, tools, KPIs) into skills, experience, and projects
        - Add realistic data analyst projects in ecommerce/logistics if missing
        - Do not exceed 1 page
        - Keep formatting, return only improved resume text ready for .docx

        --- RESUME ---
        {resume_text}

        --- JOB DESCRIPTION ---
        {jd_text}
        """

        with st.spinner("Crafting your optimized resume..."):
            client = openai.OpenAI(api_key=openai_api_key)
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional resume generator."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            improved_text = response.choices[0].message.content

        st.subheader("🔍 Preview")
        st.text_area("Modified Resume Text", improved_text, height=400)

        # Generate downloadable .docx
        def create_docx_from_text(text):
            new_doc = docx.Document()
            for line in text.split("\n"):
                if line.strip():
                    para = new_doc.add_paragraph(line.strip())
                    para.paragraph_format.space_after = Pt(3)
            return new_doc

        new_resume = create_docx_from_text(improved_text)
        bio = io.BytesIO()
        new_resume.save(bio)
        st.download_button(
            label="📥 Download Customized Resume",
            data=bio.getvalue(),
            file_name="Rafeeq_Customized_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
